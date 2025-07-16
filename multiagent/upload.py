import json
import boto3
import os
from datetime import datetime
from dotenv import load_dotenv
from botocore.exceptions import NoCredentialsError, PartialCredentialsError, ClientError
from PIL import Image
from io import BytesIO
from docx.shared import Inches
from strands import tool
import base64
# Load environment variables
# Load environment variables
load_dotenv(".env", override=True)

session = boto3.Session(
    aws_access_key_id=os.environ.get('AWS_ACCESS_KEY'),
    aws_secret_access_key=os.environ.get('AWS_SECRET_KEY'),
    region_name='ap-southeast-1'
)
s3 = session.client("s3")

        
def upload_json_to_s3(data: dict, bucket: str, prefix: str, filename: str = "output.json"):
    if hasattr(data, "dict"):
        data = data.dict()
    elif hasattr(data, "__dict__"):
        data = data.__dict__
    s3_key = f"{prefix.rstrip('/')}/{filename}"
    json_bytes = json.dumps(data, ensure_ascii=False, indent=2).encode('utf-8')
    
    s3.put_object(
        Bucket=bucket,
        Key=s3_key,
        Body=json_bytes,
        ContentType="application/json"
    )
    print(f"✅ JSON uploaded to s3://{bucket}/{s3_key}")

def read_json_from_s3(bucket = 'testworkflow123', prefix = 'info_agent/companyInfo.json'):
    try:
        response = s3.get_object(Bucket=bucket, Key=prefix)
        content = response['Body'].read().decode('utf-8')
        return json.loads(content)
    except ClientError as e:
        print(f"S3 ClientError: {e}")
    except json.JSONDecodeError as e:
        print(f"JSON decode error: {e}")
    except Exception as e:
        print(f"Unexpected error: {e}")
    return None

def upload_folder_to_s3(local_folder, bucket_name, s3_prefix=''):
    for root, dirs, files in os.walk(local_folder):
        for file in files:
            local_path = os.path.join(root, file)
            # Create relative path from the base local folder
            relative_path = os.path.relpath(local_path, local_folder)
            # Create full S3 path using the prefix
            s3_key = os.path.join(s3_prefix, relative_path).replace("\\", "/")  # for Windows

            print(f"Uploading {local_path} to s3://{bucket_name}/{s3_key}")
            s3.upload_file(local_path, bucket_name, s3_key)

def upload_text_to_s3(bucket: str, prefix: str, text: str):
    """
    Uploads a text string to an S3 bucket with a given prefix.

    Args:
        bucket (str): The name of the S3 bucket.
        prefix (str): The key prefix (simulated folder path).
        filename (str): The name of the file to upload.
        text (str): The content to upload.

    Returns:
        str: The full S3 URI of the uploaded file.
    """
    
    try:
        s3.put_object(Bucket=bucket, Key=prefix, Body=text)
        s3_uri = f"s3://{bucket}/{prefix}"
        print(f"✅ Uploaded successfully to {s3_uri}")
        return s3_uri
    except Exception as e:
        print(f"❌ Upload failed: {e}")
        return None

access = os.getenv('AWS_ACCESS_KEY')
secret = os.getenv("AWS_SECRET_KEY")


def get_image_context(image: Image.Image, chart_name="bieu_do") -> str:
    buffered = BytesIO()
    image.save(buffered, format="PNG")
    image_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")

    payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": 4096,
        "temperature": 0,
        "system": (
            "Bạn là một chuyên gia phân tích tài chính. Nhiệm vụ chính là Đưa ra nhận định và đánh giá, phân tích các thông số trong báo cáo tài chính"
        ),
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/png",
                            "data": image_base64
                        }
                    },
                    {
                        "type": "text",
                        "text": "Hãy phân tích bảng tổng quan về báo cáo tài chính này giúp tôi"
                    }
                ]
            }
        ]
    }

    session = boto3.Session(
        aws_access_key_id=access,
        aws_secret_access_key=secret,
        region_name="ap-southeast-1"
    )
    client = session.client("bedrock-runtime")

    response = client.invoke_model(
        modelId='arn:aws:bedrock:ap-southeast-1:389903776084:inference-profile/apac.anthropic.claude-3-5-sonnet-20240620-v1:0',
        contentType='application/json',
        accept='application/json',
        body=json.dumps(payload)
    )
    response_body = json.loads(response['body'].read())
    raw_text = response_body["content"][0]["text"]
    return raw_text

def download_and_insert_image_from_s3_to_paragraph(bucket_name, prefix, filename, paragraph, image_path="temp_image.png", width=Inches(3.5)):
    key = f"{prefix.rstrip('/')}/{filename}"
    print(key)
    try:
        response = s3.get_object(Bucket=bucket_name, Key=key)
        image_data = response['Body'].read()
        image = Image.open(BytesIO(image_data))
        image.save(image_path, format=image.format)
        
        run = paragraph.add_run()
        run.add_picture(image_path, width=width)
        print("✅ Ảnh đã được tải từ S3 và chèn vào ô bảng.")
        return True
    except Exception as e:
        print("❌ Lỗi tải ảnh từ S3:", repr(e))
        return False

@tool
def download_and_insert_image_from_s3(doc, bucket_name, prefix, filename, image_path="temp_image.png", width=Inches(6)):
    # s3 = boto3.client('s3')
    key = f"{prefix.rstrip('/')}/{filename}"  # Ensure no double slashes
    print(key)
    try:
        response = s3.get_object(Bucket=bucket_name, Key=key)
        image_data = response['Body'].read()
        image = Image.open(BytesIO(image_data))
        output = get_image_context(image)
        image.save(image_path, format = image.format)
        
        doc.add_paragraph(output)
        doc.add_picture(image_path, width=width)
        print("✅ Ảnh đã được tải từ S3 và chèn vào Word.")
        return True
    except Exception as e:
        print("❌ Lỗi tải ảnh từ S3:", repr(e))
        return False

if __name__ == "__main__":
    print(read_json_from_s3(bucket = 'testworkflow123', prefix = 'info_agent/companyInfo.json'))