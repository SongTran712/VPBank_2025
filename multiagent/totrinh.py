from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from PIL import Image
import requests
import json
from docx2pdf import convert
from upload import download_and_insert_image_from_s3, download_and_insert_image_from_s3_to_paragraph, read_json_from_s3, upload_json_to_s3
from strands import tool
from chartool import get_chart_context

def set_document_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_centered_title(doc, text, font_size=14):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(font_size)

def add_bold_heading(doc, text, color="000000"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(color)

def create_table(doc, rows, cols, col_widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    table.autofit = True
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = 0
    return table



# Load data from JSON
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import json
from upload import get_image_context

@tool
def get_data():
    data = read_json_from_s3('testworkflow123', 'totrinh/totrinh.json')
    if data:
        return {
            "status":'success',
            'data': data
        }
    return {
        'status':'fail'
    }

@tool
def update_to_trinh(field, input):
    valid_fields = [
        'customer', 'lich_su_phat_trien', 'ban_lanh_dao', 'danh_gia_san_pham',
        'danh_gia_thi_truong', 'danh_gia_kiem_toan', 'tinh_hinh_tai_chinh',
        'ruiro', 'ketluan'
    ]

    if field not in valid_fields:
        return {
            "message": "Field không tồn tại. Hãy chọn lại"
        }


    doc = Document()
    set_document_style(doc)
    add_centered_title(doc, "TỜ TRÌNH TÍN DỤNG KHDN")
    
    data = read_json_from_s3('testworkflow123','totrinh/totrinh.json')
    data[field] = input
    upload_json_to_s3(data, 'testworkflow123', 'totrinh', 'totrinh.json')
    
    if field == "customer":
        add_bold_heading(doc, "B. THÔNG TIN KHÁCH HÀNG", "FF0000")
        add_bold_heading(doc, "Thông tin chung về Khách hàng (KH):")
        doc.add_paragraph(input)
        
    elif field == "lich_su_phat_trien":
        add_bold_heading(doc, "C. ĐÁNH GIÁ THÔNG TIN KHÁCH HÀNG", "FF0000")
        add_bold_heading(doc, "1. Đánh giá thông tin pháp lý và bộ máy tổ chức của doanh nghiệp (BCTN):", "FF0000")
        doc.add_paragraph(input)

    elif field == "ban_lanh_dao":
        add_bold_heading(doc, "2. Cấu trúc Công Ty:", "FF0000")
        doc.add_paragraph(input)

    elif field == "danh_gia_san_pham":
        add_bold_heading(doc, "3. Đánh giá sản phẩm, dịch vụ và năng lực sản xuất, phân phối sản phẩm", "FF0000")
        doc.add_paragraph(input)

    elif field == "danh_gia_thi_truong":
        add_bold_heading(doc, "4. Đánh giá thị trường (BCTN)", "FF0000")
        doc.add_paragraph(input)

    elif field == "danh_gia_kiem_toan":
        add_bold_heading(doc, "5.1 Thông tin chung về báo cáo tài chính (BCTC):", "FF0000")
        doc.add_paragraph(input)

    elif field == "tinh_hinh_tai_chinh":
        add_bold_heading(doc, "5.2 Đánh giá tổng quan về tình hình tài chính của Doanh nghiệp:", "FF0000")
        add_bold_heading(doc, "Đánh giá chất lượng, độ tin cậy thông tin tài chính:")
        doc.add_paragraph(input)

    elif field == "ruiro":
        add_bold_heading(doc, "6. Phân tích Rủi Ro:", "FF0000")
        doc.add_paragraph(input)

    elif field == "ketluan":
        add_bold_heading(doc, "7. Kết luận:", "FF0000")
        doc.add_paragraph(input)

    # Save DOCX and convert to PDF
    doc_path = "totrinh.docx"
    pdf_path = "totrinh.pdf"
    doc.save(doc_path)
    convert(doc_path, pdf_path)

    return {
        "message": f"✅ Đã cập nhật thành công phần '{field}' trong tờ trình.",
        "doc_path": doc_path,
        "pdf_path": pdf_path
    }

@tool
def create_to_trinh(data):
    # Parse input if it's a JSON string
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError:
            raise ValueError("Input is not valid JSON")

    doc = Document()
    set_document_style(doc)

    # Title and sections
    add_centered_title(doc, "TỜ TRÌNH TÍN DỤNG KHDN")

    add_bold_heading(doc, "A. THÔNG TIN CHUNG", "FF0000")
    add_bold_heading(doc, "Mục đích Tờ trình tín dụng:")
    doc.add_paragraph("☒ Cấp tín dụng mới\n☐ Tái cấp tín dụng\n☐ Cấp tăng hạn mức")
    doc.add_paragraph("Theo các hình thức:")
    table = create_table(doc, 1, 2, [Inches(8), Inches(10)])
    table.cell(0, 0).text = "☒ Cấp hạn mức (HM) vay vốn\n☐ Cho vay theo món ngắn hạn\n☐ Cho vay theo món trung dài hạn"

    # Customer information
    add_bold_heading(doc, "\nB. THÔNG TIN KHÁCH HÀNG", "FF0000")
    add_bold_heading(doc, "Thông tin chung về Khách hàng (KH):")
    doc.add_paragraph(data.get("customer", "Không có thông tin khách hàng"))

    add_bold_heading(doc, "Tóm tắt thông tin tài chính:")
    download_and_insert_image_from_s3(doc, 'testworkflow123', 'content/fin_analyst/fin_charts/', 'Bảng tài chính tổng hơp.png')
    
    add_bold_heading(doc, "Tín dụng đã cấp cho khách hàng:")
    doc.add_picture('qhetaichinh.png', width=Inches(6))

    # # Optional: Insert image if available
    # url = data.get("image_url")
    # if url:
    #     download_and_insert_image(doc, url)



    # Customer Evaluation
    add_bold_heading(doc, "C. ĐÁNH GIÁ THÔNG TIN KHÁCH HÀNG", "FF0000")

    add_bold_heading(doc, "1. Đánh giá thông tin pháp lý và bộ máy tổ chức của doanh nghiệp (BCTN):", "FF0000")
    
    doc.add_paragraph(data.get("lich_su_phat_trien", ""))

    add_bold_heading(doc, "2. Cấu trúc Công Ty:", "FF0000")
    doc.add_paragraph(data.get("ban_lanh_dao", ""))

    add_bold_heading(doc, "3. Đánh giá sản phẩm, dịch vụ và năng lực sản xuất, phân phối sản phẩm", "FF0000")
    doc.add_paragraph(data.get("danh_gia_san_pham", ""))

    add_bold_heading(doc, "4. Đánh giá thị trường (BCTN)", "FF0000")
    doc.add_paragraph(data.get("danh_gia_thi_truong", ""))

    add_bold_heading(doc, "5. Đánh giá tình hình tài chính", "FF0000")

    add_bold_heading(doc, "5.1 Thông tin chung về báo cáo tài chính (BCTC):")
    doc.add_paragraph(data.get("danh_gia_kiem_toan", ""))

    add_bold_heading(doc, "5.2 Đánh giá tổng quan về tình hình tài chính của Doanh nghiệp:")
    add_bold_heading(doc, "Đánh giá chất lượng, độ tin cậy thông tin tài chính:")
    doc.add_paragraph(data.get("tinh_hinh_tai_chinh", ""))
    
    add_bold_heading(doc, "Chart Cơ Cấu Vốn, Cơ Cấu Tài Sản Nợ")
    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    download_and_insert_image_from_s3_to_paragraph(
    'testworkflow123', 'content/fin_analyst/fin_charts/', 'Cơ cấu vốn.png',
    paragraph=table.cell(0, 0).paragraphs[0],
    width=Inches(3)
)

    download_and_insert_image_from_s3_to_paragraph(
    'testworkflow123', 'content/fin_analyst/fin_charts/', 'Cơ cấu tài sản nợ.png',
    paragraph=table.cell(0, 1).paragraphs[0],
    width=Inches(3)
)
    add_bold_heading(doc, "Chart Biên lợi nhuận và lợi nhuận, Phân tích lợi nhuận")

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    download_and_insert_image_from_s3_to_paragraph(
    'testworkflow123', 'content/fin_analyst/fin_charts/', 'Biên lợi nhuận và lợi nhuận.png',
    paragraph=table.cell(0, 0).paragraphs[0],
    width=Inches(3)
)

    download_and_insert_image_from_s3_to_paragraph(
    'testworkflow123', 'content/fin_analyst/fin_charts/', 'Phân tích lợi nhuận.png',
    paragraph=table.cell(0, 1).paragraphs[0],
    width=Inches(3)
)
    
    add_bold_heading(doc, "Tài sản và hiệu quả sinh lời, Vốn chủ sở hữu và khả năng sinh lời")

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Inches(3.5)
    table.columns[1].width = Inches(3.5)

    download_and_insert_image_from_s3_to_paragraph(
    'testworkflow123', 'content/fin_analyst/fin_charts/', 'Tài sản và hiệu quả sinh lời.png',
    paragraph=table.cell(0, 0).paragraphs[0],
    width=Inches(3)
)

    download_and_insert_image_from_s3_to_paragraph(
    'testworkflow123', 'content/fin_analyst/fin_charts/', 'Vốn chủ sở hữu và khả năng sinh lời.png',
    paragraph=table.cell(0, 1).paragraphs[0],
    width=Inches(3)
)    
    
    add_bold_heading(doc, "6. Phân tích Rủi Ro:",  "FF0000")
    # add_bold_heading(doc, "Đánh giá chất lượng, độ tin cậy thông tin tài chính:")
    doc.add_paragraph(data.get("ruiro", ""))
    
    add_bold_heading(doc,"7. Kết luận: ", "FF0000")
    doc.add_paragraph(data.get("ketluan",""))
    
    # Save document
    doc_path = "totrinh.docx"
    pdf_path = "totrinh.pdf"
    doc.save(doc_path)
    convert(doc_path, pdf_path)
    
    totrinh = {
  "customer": data.get("customer", ""),
  "lich_su_phat_trien": data.get("lich_su_phat_trien", ""),
  "ban_lanh_dao": data.get("ban_lanh_dao", ""),
  "danh_gia_thi_truong": data.get("danh_gia_thi_truong", "Đánh giá tình hình thị trường mà công ty đang hoạt động."),
  "danh_gia_san_pham": data.get(""),
  "danh_gia_kiem_toan": data.get("kiem_toan", ""),
  "tinh_hinh_tai_chinh": data.get("tinh_hinh_tai_chinh", ""),
  "ruiro": data.get("ruiro", ""),
  "ketluan": data.get("ketluan", "")
}
    upload_json_to_s3(totrinh, 'testworkflow123', 'totrinh', 'totrinh.json')
    return {
        "message": "✅ Đã tạo thành công tờ trình tín dụng.",
        "doc_path": doc_path,
        "pdf_path": pdf_path
    }

if __name__ == "__main__":
#     create_to_trinh({
#     "customer": "CÔNG TY CỔ PHẦN VIMC LOGISTICS\nWebsite: Không có thông tin\nSố điện thoại: 04-35772047/48\nEmail: info@vimclogistics.vn\nMã số thuế: 0102345275\nNgười đại diện theo pháp luật: Bà Đinh Thị Việt Hà",
#     "lich_su_phat_trien": "Công ty VIMC Logistics được thành lập với mục đích cung cấp các dịch vụ logistics toàn diện, bao gồm vận tải đường bộ, đường sắt, và đường biển, cũng như dịch vụ xuất nhập khẩu và môi giới thương mại. Hiện nay, công ty không có thông tin cụ thể về tập đoàn mẹ hoặc cổ đông lớn.",
#     "ban_lanh_dao": "Cấu trúc tổ chức của công ty bao gồm:\n- Chủ tịch HĐQT: Ông Mai Lê Lợi\n- Thành viên HĐQT: Bà Dương Thu Hiền\n- Thành viên HĐQT: Ông Phan Nhân Hải\n- Thành viên HĐQT: Ông Nguyễn Quốc Cường\n- Thành viên HĐQT: Bà Đinh Thị Việt Hà\n- Phụ trách quản trị: Bà Vũ Thị Thanh Nhàn\n- Tổng Giám đốc: Bà Đinh Thị Việt Hà\n- Phó Tổng Giám đốc: Ông Phạm Bá Ngân",
#     "danh_gia_thi_truong": "Công ty VIMC Logistics hoạt động trong ngành logistics và vận tải, một lĩnh vực đang phát triển mạnh mẽ do nhu cầu giao thông vận tải ngày càng tăng. Thị trường này có nhiều cơ hội tăng trưởng nhưng cũng đối mặt với cạnh tranh cao từ các công ty lớn và nhỏ.",
#     "danh_gia_san_pham": "Công ty cung cấp dịch vụ logistics toàn diện, bao gồm vận tải đường bộ, đường sắt, và đường biển, cũng như dịch vụ xuất nhập khẩu và môi giới thương mại. Sản phẩm và dịch vụ của công ty có tính cạnh tranh cao do đội ngũ quản lý và nhân viên có kinh nghiệm, cùng với mạng lưới đối tác mạnh mẽ.",
#     "danh_gia_kiem_toan": "Tình trạng kiểm toán báo cáo tài chính của công ty hiện nay không có thông tin cụ thể về đơn vị kiểm toán và nội dung nổi bật của báo cáo.",
#     "tong_quan_tai_chinh": "Tổng quan tình hình tài chính của công ty VIMC Logistics hiện nay không có thông tin cụ thể. Cần tiếp tục theo dõi và cập nhật báo cáo tài chính để đánh giá chính xác."
# }
# )
    print(update_to_trinh('customer','hihi'))

